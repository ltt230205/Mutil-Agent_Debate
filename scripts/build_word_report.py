"""Build the Vietnamese academic report draft for Chapters 1 and 2."""

from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = ROOT / "report" / "Bao_cao_Multi_Agent_Debate_Chuong_1_2.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "17365D"
LIGHT_BLUE = "DCE6F1"
PALE_BLUE = "EDF3F8"
LIGHT_GRAY = "F4F6F9"
MID_GRAY = "D9E2F3"
DARK_GRAY = "404040"
WHITE = "FFFFFF"


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_twips: int = 9360) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_twips))
    tbl_w.set(qn("w:type"), "dxa")
    table.autofit = False


def add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    run_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_pr.extend([color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([run_pr, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_equation(document: Document, expression: str, label: str | None = None) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    math_para = OxmlElement("m:oMathPara")
    math_para_pr = OxmlElement("m:oMathParaPr")
    justification = OxmlElement("m:jc")
    justification.set(qn("m:val"), "center")
    math_para_pr.append(justification)
    math_para.append(math_para_pr)
    math = OxmlElement("m:oMath")
    math_run = OxmlElement("m:r")
    math_run_pr = OxmlElement("m:rPr")
    math_style = OxmlElement("m:sty")
    math_style.set(qn("m:val"), "p")
    math_run_pr.append(math_style)
    math_text = OxmlElement("m:t")
    math_text.text = expression
    math_run.extend([math_run_pr, math_text])
    math.append(math_run)
    math_para.append(math)
    paragraph._p.append(math_para)
    if label:
        label_run = paragraph.add_run(f"   {label}")
        label_run.font.name = "Cambria Math"
        label_run.font.size = Pt(10)


def add_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True


def add_note(document: Document, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    set_table_width(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, PALE_BLUE)
    set_cell_margins(cell, 120, 160, 120, 160)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullets(document: Document, items: Iterable[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_numbered(document: Document, items: Iterable[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Number")


def add_table(
    document: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    widths: Sequence[int] | None = None,
    font_size: float = 9.0,
) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    if widths is None:
        base = 9360 // len(headers)
        widths = [base] * len(headers)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, LIGHT_GRAY)
        set_cell_width(cell, widths[idx])
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor.from_string(NAVY)
    set_repeat_table_header(table.rows[0])
    for row_values in rows:
        row = table.add_row()
        prevent_row_split(row)
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            run = p.add_run(value)
            run.font.size = Pt(font_size)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_architecture_figure(document: Document) -> None:
    stages = [
        ("CÂU HỎI ĐẦU VÀO", "context + question + options + sample_id"),
        ("SOLVER ĐỘC LẬP", "N đáp án ban đầu, rationale tóm tắt, evidence, confidence"),
        ("PHÁT HIỆN BẤT ĐỒNG", "so sánh nhãn đáp án, mức đồng thuận và trace"),
        ("PHẢN BIỆN CHUYÊN BIỆT", "Critic + Skeptic + Evidence Checker"),
        ("REVISION", "mỗi Solver sửa hoặc bảo lưu đáp án có giải trình"),
        ("LẶP THEO SỐ VÒNG", "r = 0, 1, 2, 3; dừng sớm nếu thỏa điều kiện"),
        ("QUYẾT ĐỊNH", "Majority Vote hoặc Blind Evidence-Aware Judge"),
        ("ĐẦU RA VÀ NHẬT KÝ", "final answer + confidence + token + latency + raw JSONL"),
    ]
    table = document.add_table(rows=len(stages) * 2 - 1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 7920)
    for idx, (title, subtitle) in enumerate(stages):
        row = table.rows[idx * 2]
        prevent_row_split(row)
        cell = row.cells[0]
        set_cell_width(cell, 7920)
        set_cell_margins(cell, 120, 180, 120, 180)
        shade_cell(cell, LIGHT_BLUE if idx % 2 == 0 else PALE_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(title)
        r1.bold = True
        r1.font.color.rgb = RGBColor.from_string(NAVY)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(subtitle)
        r2.font.size = Pt(9)
        if idx < len(stages) - 1:
            arrow_cell = table.rows[idx * 2 + 1].cells[0]
            arrow_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            arrow_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            arrow_run = arrow_cell.paragraphs[0].add_run("↓")
            arrow_run.bold = True
            arrow_run.font.size = Pt(12)
            arrow_run.font.color.rgb = RGBColor.from_string(BLUE)
    add_caption(document, "Hình 1. Kiến trúc tổng thể của hệ thống Multi-Agent Debate đề xuất")


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("202020")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(BLUE)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string(BLUE)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(10)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor.from_string(DARK_GRAY)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)

    if "Reference" not in styles:
        ref_style = styles.add_style("Reference", WD_STYLE_TYPE.PARAGRAPH)
    else:
        ref_style = styles["Reference"]
    ref_style.font.name = "Calibri"
    ref_style.font.size = Pt(9.5)
    ref_style.paragraph_format.left_indent = Inches(0.32)
    ref_style.paragraph_format.first_line_indent = Inches(-0.32)
    ref_style.paragraph_format.space_after = Pt(5)
    ref_style.paragraph_format.line_spacing = 1.05

    if "Code Block" not in styles:
        code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = styles["Code Block"]
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(8.5)
    code_style.paragraph_format.left_indent = Inches(0.2)
    code_style.paragraph_format.right_indent = Inches(0.2)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(4)
    code_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def configure_section(section, content: bool = True) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    if content:
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("BEYOND MAJORITY VOTING | BẢN THẢO CHƯƠNG 1–2")
        run.font.name = "Calibri"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("7F7F7F")
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        add_field(p, "PAGE", "1")
        pg_num_type = OxmlElement("w:pgNumType")
        pg_num_type.set(qn("w:start"), "1")
        section._sectPr.append(pg_num_type)


def add_cover(document: Document) -> None:
    section = document.sections[0]
    configure_section(section, content=False)
    section.different_first_page_header_footer = True

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(74)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BÁO CÁO TIỂU LUẬN KHOA HỌC")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("BEYOND MAJORITY VOTING")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor.from_string(NAVY)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(26)
    r = p.add_run("Multi-Agent Debate for Reliable NLP Reasoning")
    r.font.name = "Calibri"
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor.from_string(BLUE)

    rule = document.add_table(rows=1, cols=1)
    rule.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(rule, 4200)
    cell = rule.cell(0, 0)
    set_cell_width(cell, 4200)
    shade_cell(cell, BLUE)
    cell.height = Pt(4)
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(26)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("BẢN THẢO HỌC THUẬT")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(DARK_GRAY)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Hoàn thiện: Mở đầu, Chương 1 và Chương 2")
    r.font.size = Pt(11)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(92)
    r = p.add_run("Chương 3, Chương 4 và kết luận thực nghiệm sẽ được bổ sung sau khi chạy main experiment")
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string("666666")

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Ngôn ngữ: Tiếng Việt | Chuẩn trích dẫn: IEEE")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(DARK_GRAY)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("2026")
    r.bold = True
    r.font.size = Pt(12)


def add_front_matter(document: Document) -> None:
    section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, content=True)

    document.add_heading("THÔNG TIN VÀ PHẠM VI BẢN THẢO", level=1)
    add_note(
        document,
        "Tài liệu này hoàn thiện phần Mở đầu, Chương 1 và Chương 2. Những phát biểu về hiệu quả thực nghiệm, "
        "xếp hạng phương pháp và trạng thái các giả thuyết chưa được đưa ra. Chương 3, Chương 4 và kết luận khoa học "
        "sẽ chỉ được hoàn thiện từ raw output của API thật sau khi pilot, main experiment và ablation kết thúc.",
    )
    add_table(
        document,
        ["Thuộc tính", "Giá trị"],
        [
            ("Tên đề tài", "Beyond Majority Voting: Multi-Agent Debate for Reliable NLP Reasoning"),
            ("Loại tài liệu", "Báo cáo tiểu luận khoa học – bản thảo Chương 1 và Chương 2"),
            ("Ngôn ngữ", "Tiếng Việt có dấu; tên biến, trường JSON và mã nguồn bằng tiếng Anh"),
            ("Chuẩn trích dẫn", "IEEE"),
            ("Đối tượng nghiên cứu", "Suy luận logic và suy luận tri thức thường thức bằng mô hình ngôn ngữ lớn"),
            ("Nguyên tắc dữ liệu", "Không tạo số liệu giả; không dùng kết quả mock/smoke để kết luận khoa học"),
        ],
        widths=[2200, 7160],
    )

    document.add_heading("MỤC LỤC", level=1)
    toc = document.add_paragraph()
    add_field(toc, 'TOC \\o "1-3" \\h \\z \\u', "Nhấn F9 trong Microsoft Word để cập nhật mục lục.")
    toc.paragraph_format.space_after = Pt(8)
    document.add_page_break()

    document.add_heading("DANH MỤC TỪ VIẾT TẮT", level=1)
    add_table(
        document,
        ["Viết tắt", "Thuật ngữ", "Diễn giải sử dụng trong báo cáo"],
        [
            ("LLM", "Large Language Model", "Mô hình ngôn ngữ lớn"),
            ("NLP", "Natural Language Processing", "Xử lý ngôn ngữ tự nhiên"),
            ("CoT", "Chain-of-Thought", "Chuỗi lập luận trung gian được mô hình chủ động xuất ra"),
            ("SC", "Self-Consistency", "Lấy mẫu nhiều đường suy luận độc lập và bỏ phiếu đáp án"),
            ("MAD", "Multi-Agent Debate", "Tranh luận nhiều tác tử dựa trên mô hình ngôn ngữ"),
            ("RQ", "Research Question", "Câu hỏi nghiên cứu"),
            ("ECE", "Expected Calibration Error", "Sai số hiệu chuẩn kỳ vọng"),
            ("CI", "Confidence Interval", "Khoảng tin cậy"),
            ("JSONL", "JSON Lines", "Định dạng một đối tượng JSON trên mỗi dòng"),
        ],
        widths=[1200, 2600, 5560],
    )


def add_introduction(document: Document) -> None:
    heading = document.add_heading("MỞ ĐẦU", level=1)
    heading.paragraph_format.page_break_before = True

    document.add_heading("1. Lý do chọn đề tài", level=2)
    document.add_paragraph(
        "Mô hình ngôn ngữ lớn đã trở thành hạ tầng cốt lõi cho nhiều hệ thống hỏi–đáp, trợ lý tri thức, hỗ trợ ra quyết định và tự động hóa xử lý văn bản. Kiến trúc Transformer [1] và khả năng học trong ngữ cảnh ở quy mô lớn [2] tạo ra bước tiến đáng kể về độ trôi chảy và khả năng thích nghi nhiệm vụ. Tuy nhiên, chất lượng ngôn ngữ bề mặt không đồng nghĩa với suy luận đáng tin cậy. Một câu trả lời có thể được diễn đạt thuyết phục nhưng dựa trên tiền đề sai, bỏ sót điều kiện, sử dụng tri thức không có trong đề hoặc thể hiện confidence không tương xứng với xác suất đúng. Khoảng cách giữa tính thuyết phục và tính đúng đắn đặc biệt nguy hiểm trong các bài toán logical reasoning và commonsense reasoning, nơi chỉ một sai lệch nhỏ trong diễn giải có thể thay đổi toàn bộ đáp án."
    )
    document.add_paragraph(
        "Các kỹ thuật inference-time như Chain-of-Thought (CoT) [3], zero-shot reasoning [4] và Self-Consistency [5] cho thấy hiệu năng có thể được cải thiện mà không cần huấn luyện lại mô hình. Trong số đó, Majority Voting là cơ chế tổng hợp phổ biến: hệ thống tạo nhiều dự đoán rồi chọn nhãn nhận nhiều phiếu nhất. Cách làm này rẻ về mặt thiết kế và dễ tái lập, nhưng chỉ khai thác tần suất đáp án, không đánh giá chất lượng bằng chứng. Nếu các mẫu suy luận mắc lỗi tương quan, số đông có thể làm một sai lầm trở nên ổn định hơn thay vì sửa nó."
    )
    document.add_paragraph(
        "Multi-Agent Debate mở rộng ensemble tĩnh thành một giao thức giao tiếp: các tác tử giải độc lập, chỉ ra lỗi, tìm phản ví dụ, kiểm tra bằng chứng và sửa đáp án qua nhiều vòng. Một số nghiên cứu gần đây báo cáo lợi ích của debate đối với factuality, reasoning và khả năng thuyết phục [6]–[8]. Dẫu vậy, debate dùng nhiều lượt gọi, nhiều token và độ trễ cao hơn. Nếu không kiểm soát inference budget, chưa thể biết phần tăng độ chính xác đến từ giao tiếp hay chỉ từ lượng compute lớn hơn. Đây là lý do nghiên cứu cần đi “beyond majority voting” nhưng đồng thời phải xem Majority Voting là đối chứng nghiêm túc, không phải một baseline yếu được dựng để phương pháp đề xuất dễ thắng."
    )

    document.add_heading("2. Bài toán nghiên cứu", level=2)
    document.add_paragraph(
        "Nghiên cứu xem xét bài toán trắc nghiệm suy luận với mỗi mẫu gồm ngữ cảnh, câu hỏi, tập lựa chọn và đáp án chuẩn. Hệ thống phải trả về một nhãn đáp án, rationale tóm tắt có thể quan sát, bằng chứng đã sử dụng và confidence trong khoảng từ 0 đến 1. Đối tượng so sánh gồm tác tử đơn, Self-Consistency, nhiều tác tử độc lập bỏ phiếu và nhiều biến thể debate. Trọng tâm không chỉ là Accuracy cuối cùng mà còn là quá trình chuyển trạng thái: khi nào một đáp án sai được sửa, khi nào đáp án đúng bị làm hỏng, mức bất đồng có mang tính chiến lược hay chỉ là khác biệt diễn đạt, và cái giá phải trả bằng token cũng như thời gian."
    )
    document.add_paragraph(
        "Độ tin cậy trong báo cáo được hiểu theo nghĩa đa chiều. Một phương pháp đáng tin không nhất thiết là phương pháp có Accuracy cao nhất trong một lần chạy; nó cần duy trì kết quả qua nhiều seed, có khoảng tin cậy hợp lý, thể hiện calibration tốt, kiểm soát được failure mode và cho phép truy vết quyết định. Vì vậy, đơn vị phân tích là cả dự đoán cuối, tập reasoning trace, transcript phản biện, token usage, latency và sự chuyển đổi đúng–sai qua các vòng."
    )

    document.add_heading("3. Mục tiêu nghiên cứu", level=2)
    document.add_paragraph("Mục tiêu tổng quát là đánh giá liệu giao thức Multi-Agent Debate có cải thiện độ tin cậy của NLP reasoning so với Majority Voting dưới điều kiện so sánh công bằng về compute hay không.")
    add_bullets(
        document,
        [
            "Đo sự đa dạng ở ba tầng: bất đồng đáp án, khác biệt ngữ nghĩa giữa rationale và khác biệt chiến lược giải.",
            "So sánh Accuracy, calibration, token cost và latency giữa các baseline và biến thể debate.",
            "Xác định số vòng debate phù hợp thông qua Correction Rate, Degradation Rate và đường cong chi phí–hiệu năng.",
            "Đánh giá lợi ích biên của role specialization so với các agent đồng nhất.",
            "Xây dựng taxonomy để phân tích successful correction, resistant error, harmful revision và lỗi của Judge.",
            "Bảo đảm khả năng tái lập bằng cấu hình YAML, seed cố định, sample ID khóa trước, cache và raw JSONL.",
        ],
    )

    document.add_heading("4. Câu hỏi nghiên cứu", level=2)
    add_table(
        document,
        ["Mã", "Câu hỏi nghiên cứu"],
        [
            ("RQ1", "MAD có tạo reasoning diversity cao hơn Self-Consistency và Multi-Agent Majority Voting hay không?"),
            ("RQ2", "Khi inference budget tương đương, MAD có cải thiện Accuracy và độ tin cậy so với Majority Voting hay không?"),
            ("RQ3", "Số vòng debate ảnh hưởng thế nào đến Accuracy, Correction Rate, Degradation Rate, token cost, latency và diversity?"),
            ("RQ4", "Các agent chuyên biệt có hiệu quả hơn các agent đồng nhất dùng cùng prompt hay không?"),
            ("RQ5", "Trong trường hợp nào debate sửa được lỗi, không sửa được lỗi hoặc làm đáp án đúng trở thành sai?"),
        ],
        widths=[900, 8460],
    )

    document.add_heading("5. Giả thuyết nghiên cứu", level=2)
    add_bullets(
        document,
        [
            "H1: MAD tạo semantic reasoning diversity cao hơn Multi-Agent Majority Voting.",
            "H2: MAD đạt Accuracy cao hơn Majority Voting, đặc biệt trên các mẫu có bất đồng ban đầu cao.",
            "H3: Agent chuyên biệt đạt hiệu năng tốt hơn agent đồng nhất.",
            "H4: Accuracy không tăng tuyến tính theo số vòng; sau một ngưỡng, hiệu quả bão hòa hoặc giảm trong khi token cost và latency tiếp tục tăng.",
            "H5: Debate sửa được một phần đáp án sai nhưng đồng thời tạo Degradation Rate do conformity pressure, sycophancy, error propagation và Judge bias.",
        ],
    )
    add_note(
        document,
        "Các giả thuyết trên được giữ nguyên trước main experiment. Sau thực nghiệm, mỗi giả thuyết chỉ được gán một trong ba trạng thái: được hỗ trợ, được hỗ trợ một phần hoặc không được hỗ trợ.",
    )

    document.add_heading("6. Đối tượng và phạm vi nghiên cứu", level=2)
    document.add_paragraph(
        "Nghiên cứu tập trung vào suy luận trắc nghiệm bằng LLM trên hai miền bổ sung nhau: logical reasoning và commonsense reasoning. LogiQA [12] được ưu tiên cho các câu hỏi đòi hỏi liên kết điều kiện và loại trừ logic; CommonsenseQA [13] được dùng cho tri thức thường thức với các phương án nhiễu gần nghĩa. StrategyQA [14] và GSM8K [15] được xem là benchmark mở rộng nếu ngân sách cho phép, không phải dữ liệu chính của thiết kế hiện tại. Phạm vi không bao gồm fine-tuning, truy cập hidden chain-of-thought, huấn luyện reward model hoặc đánh giá an toàn ở miền rủi ro cao."
    )
    document.add_paragraph(
        "Mỗi dataset dự kiến dùng 200–300 mẫu trên tập đánh giá đã khóa, ba seed 42, 123 và 2026, năm solver mặc định, và số vòng 0, 1, 2, 3. Prompt được phát triển trên development set; test hoặc test subset đã khóa không được dùng để điều chỉnh prompt. API key chỉ được đọc từ biến môi trường, model name nằm trong YAML, và mọi phản hồi model phải được lưu để có thể đánh giá lại mà không gọi API lần nữa."
    )

    document.add_heading("7. Phương pháp nghiên cứu", level=2)
    document.add_paragraph(
        "Nghiên cứu sử dụng thiết kế thực nghiệm đối chứng theo từng mẫu. Bốn baseline tối thiểu gồm Single Direct, Single CoT, Self-Consistency và Multi-Agent Majority Voting. Nhóm phương pháp debate gồm Homogeneous Debate, Role-Specialized Debate + Majority, Role-Specialized Debate + Judge và Evidence-Aware Judge. Fair-compute được triển khai theo hai trục: khớp số model call và khớp tổng token. Cách so sánh ghép cặp cho phép áp dụng McNemar test [19] trên các cặp dự đoán đúng–sai của cùng mẫu; bootstrap [20] được dùng để ước lượng khoảng tin cậy 95%."
    )
    document.add_paragraph(
        "Phân tích định lượng được bổ sung bằng behavioral analysis và error analysis. Các nhóm hành vi bao gồm Successful Correction, Resistant Error, Harmful Revision, Productive Disagreement và Minority-Correct Case. Taxonomy lỗi bao phủ lỗi logic, diễn giải sai, thiếu bằng chứng, giả định không được hỗ trợ, lỗi số học, hallucination, conformity, lỗi Judge, lỗi trích xuất đáp án và context overload. Khi cần gán nhãn thủ công, mức đồng thuận giữa người gán nhãn có thể báo cáo bằng Cohen's kappa [24]."
    )

    document.add_heading("8. Đóng góp dự kiến", level=2)
    add_numbered(
        document,
        [
            "Một giao thức Multi-Agent Debate tách rõ communication protocol khỏi decision protocol.",
            "Một thiết kế fair-compute so sánh debate và voting theo model call, token budget và Accuracy trên mỗi 1.000 token.",
            "Một bộ đo reasoning diversity không đồng nhất sự khác biệt câu chữ với sự khác biệt chiến lược.",
            "Một quy trình truy vết từng claim bằng nhãn SUPPORTED, UNSUPPORTED, CONTRADICTED và UNCERTAIN.",
            "Một pipeline tái lập với JSON contract, retry hữu hạn, cache, checkpoint, raw JSONL, seed control và cấu hình ngoài mã nguồn.",
            "Một khuôn phân tích trung thực, báo cáo cả các trường hợp debate thua Majority Voting và các harmful revisions.",
        ],
    )

    document.add_heading("9. Cấu trúc báo cáo", level=2)
    document.add_paragraph(
        "Sau phần Mở đầu, Chương 1 trình bày cơ sở lý thuyết về LLM reasoning, CoT, Self-Consistency, hệ đa tác tử, Majority Voting, debate và độ tin cậy. Chương 2 đặc tả phương pháp đề xuất, vai trò, giao thức giao tiếp, giao thức quyết định, hợp đồng JSON, kiểm soát chi phí và cơ chế hạn chế failure mode. Chương 3 sẽ mô tả dữ liệu, cấu hình mô hình, baseline, fair-compute, ablation và quy trình đánh giá sau khi cấu hình thực nghiệm được khóa. Chương 4 chỉ được viết từ kết quả API thật, gồm định lượng, kiểm định thống kê, behavioral analysis, error analysis, trả lời RQ và kiểm định H1–H5. Phần cuối tổng kết, nêu hạn chế và hướng phát triển dựa trên bằng chứng thu được."
    )

    document.add_heading("10. Ánh xạ RQ, giả thuyết, thực nghiệm và chỉ số", level=2)
    add_caption(document, "Bảng 1. Ma trận truy vết thiết kế nghiên cứu")
    add_table(
        document,
        ["RQ", "Giả thuyết", "Thực nghiệm chính", "Chỉ số/kiểm định"],
        [
            ("RQ1", "H1", "B3/B4 so với debate đồng nhất và chuyên biệt", "Disagreement; semantic diversity; strategy labels; bootstrap CI"),
            ("RQ2", "H2", "Fair-compute theo số call và tổng token", "Accuracy; McNemar; win/loss/tie; Accuracy/1.000 token"),
            ("RQ3", "H4", "Ablation số vòng 0, 1, 2, 3", "Accuracy; correction; degradation; token; mean/median/P95 latency"),
            ("RQ4", "H3", "Homogeneous vs specialized; remove-one-role", "Accuracy; diversity; chi phí; error taxonomy"),
            ("RQ5", "H5", "Behavioral transition và case study có đối chứng", "Successful/resistant/harmful; minority-correct; Judge error"),
        ],
        widths=[700, 900, 3300, 4460],
        font_size=8.5,
    )


def add_chapter_1(document: Document) -> None:
    heading = document.add_heading("CHƯƠNG 1. TỔNG QUAN VÀ CƠ SỞ LÝ THUYẾT", level=1)
    heading.paragraph_format.page_break_before = True

    document.add_heading("1.1. Large Language Models và NLP Reasoning", level=2)
    document.add_paragraph(
        "LLM hiện đại phần lớn dựa trên Transformer, kiến trúc sử dụng attention để mô hình hóa quan hệ giữa các token mà không cần xử lý tuần tự như mạng hồi tiếp [1]. Khi quy mô dữ liệu, số tham số và năng lực học trong ngữ cảnh tăng, mô hình có thể thực hiện nhiều nhiệm vụ chỉ từ hướng dẫn hoặc một số ví dụ trong prompt [2]. Thành tựu này khiến một mô hình sinh tổng quát có thể được dùng như solver, critic hoặc judge chỉ bằng cách thay đổi vai trò và context. Tuy nhiên, việc dùng chung một mô hình nền cũng tạo ra nguy cơ correlated error: các agent bề ngoài độc lập vẫn chia sẻ tham số, dữ liệu huấn luyện và thiên kiến giải mã."
    )
    document.add_paragraph(
        "Trong báo cáo này, NLP reasoning là khả năng kết hợp thông tin để tạo ra quyết định mà không thể giải thích đầy đủ bằng khớp từ khóa trực tiếp. Logical reasoning đòi hỏi tôn trọng quan hệ kéo theo, phủ định, thứ tự, ràng buộc và lượng từ trong ngữ cảnh. Commonsense reasoning yêu cầu kích hoạt tri thức thường thức phù hợp, đồng thời phân biệt tri thức hợp lý với giả định tùy tiện. Cả hai đều có thể được đánh giá dưới dạng multiple choice, nhưng nguồn lỗi khác nhau: LogiQA thường nhạy với việc bỏ sót điều kiện, còn CommonsenseQA dễ bị tác động bởi liên tưởng bề mặt và độ phổ biến từ vựng."
    )
    document.add_paragraph(
        "Một distinction quan trọng là đáp án, rationale quan sát được và cơ chế nội tại của mô hình không phải cùng một đối tượng. Hệ thống chỉ thu thập structured rationale mà mô hình chủ động sinh ra; không tuyên bố truy cập hay khôi phục hidden chain-of-thought. Rationale được dùng như một artifact để kiểm tra claim, so sánh chiến lược và hỗ trợ quyết định, nhưng không được xem là bằng chứng tuyệt đối về quá trình tính toán nội tại. Cách tiếp cận này vừa phù hợp với giới hạn kỹ thuật, vừa tránh diễn giải quá mức nội dung do mô hình tạo."
    )
    document.add_paragraph(
        "Đánh giá reasoning vì thế cần hai tầng. Tầng task-level đo câu trả lời có đúng hay không. Tầng process-level xem xét bằng chứng được viện dẫn, tính nhất quán của các bước, khả năng phản hồi phê bình và pattern lỗi. Một phương pháp có thể tăng Accuracy nhưng giảm khả năng truy vết, hoặc tạo rationale đa dạng nhưng không cải thiện correctness. Thiết kế nghiên cứu phải giữ các đại lượng này tách biệt để không biến một tín hiệu phụ thành kết luận chính."
    )

    document.add_heading("1.2. Chain-of-Thought và Self-Consistency", level=2)
    document.add_heading("1.2.1. Chain-of-Thought có thể quan sát", level=3)
    document.add_paragraph(
        "Chain-of-Thought prompting yêu cầu mô hình tạo các bước trung gian trước khi đưa ra đáp án, qua đó có thể cải thiện hiệu năng trên nhiều bài toán số học, symbolic reasoning và commonsense [3]. Zero-shot CoT cho thấy chỉ một chỉ dẫn khuyến khích suy nghĩ theo bước cũng có thể kích hoạt hành vi suy luận ở mô hình đủ lớn [4]. Về mặt hệ thống, CoT tạo ra một bề mặt để agent khác kiểm tra: Critic có thể định vị bước suy ra sai, Skeptic có thể đặt phản ví dụ, còn Evidence Checker có thể đối chiếu claim với premise."
    )
    document.add_paragraph(
        "Dù vậy, rationale dài không mặc nhiên tốt hơn. Một trace có thể hợp lý hóa hậu nghiệm cho đáp án đã chọn, chèn chi tiết không cần thiết hoặc khiến Judge thiên vị vì độ trôi chảy. Vì vậy, Solver trong nghiên cứu chỉ sinh rationale_summary ngắn, đánh số ý và liệt kê evidence. Decision prompt cấm Judge ưu tiên một câu trả lời chỉ vì dài hơn. Biện pháp này nhằm giảm verbosity bias đồng thời giữ đủ thông tin để audit."
    )
    document.add_heading("1.2.2. Self-Consistency", level=3)
    document.add_paragraph(
        "Self-Consistency thay thế một đường giải duy nhất bằng K lần lấy mẫu độc lập, sau đó chọn đáp án phổ biến nhất [5]. Trực giác là một bài toán có thể có nhiều đường suy luận hợp lệ, trong khi lỗi ngẫu nhiên phân tán giữa các mẫu. Nếu các lỗi gần độc lập và xác suất một mẫu đúng lớn hơn ngẫu nhiên, tổng hợp nhiều mẫu có thể giảm phương sai. Với tập đáp án C và các dự đoán y₁,…,yₖ, luật quyết định phổ biến là:"
    )
    add_equation(document, "ŷSC = arg max(y ∈ C) Σᵢ₌₁ᴷ 𝟙(yᵢ = y)", "(1)")
    document.add_paragraph(
        "Giả định “lỗi đủ độc lập” là điểm dễ bị bỏ qua. Khi mọi sample dùng cùng model, cùng prompt và cùng context, chúng có thể hội tụ vào cùng shortcut sai. Temperature cao hơn có thể tăng lexical diversity nhưng cũng tăng noise; temperature thấp tạo ổn định nhưng giảm khám phá. Do đó, nghiên cứu không chỉ báo cáo K mà còn giữ nguyên model, temperature, max output tokens và seed policy khi so sánh. Self-Consistency là đối chứng đặc biệt quan trọng cho fair-compute bởi nó cũng tiêu thụ nhiều lượt gọi nhưng không có giao tiếp."
    )

    document.add_heading("1.3. Hệ thống đa tác tử sử dụng LLM", level=2)
    document.add_paragraph(
        "Một hệ đa tác tử LLM có thể được mô tả bởi bốn thành phần: tập agent, trạng thái dùng chung, giao thức truyền thông và quy tắc ra quyết định. Agent có thể là nhiều mô hình khác nhau hoặc nhiều instance của cùng mô hình với role prompt khác nhau. Trạng thái dùng chung chứa câu hỏi, câu trả lời trước, critique và evidence. Giao thức quy định agent nào được đọc thông tin nào ở vòng nào; quy tắc quyết định chuyển tập phát biểu thành một đáp án cuối. Việc tách bốn thành phần giúp xác định đúng nguyên nhân nếu hiệu năng thay đổi."
    )
    document.add_paragraph(
        "Các phương pháp self-improvement như Self-Refine [9] và Reflexion [10] cho thấy feedback ngôn ngữ có thể hướng dẫn sửa đầu ra mà không cập nhật trọng số. ReAct [11] mở rộng reasoning bằng cách xen kẽ suy luận với hành động và quan sát. Debate nằm trong cùng họ inference-time interaction nhưng khác ở nguồn feedback: thay vì một agent tự phản tỉnh, nhiều vai trò có mục tiêu kiểm tra khác nhau cùng tác động lên revision. Lợi ích kỳ vọng là phân tách trách nhiệm nhận thức; rủi ro là thông tin sai từ một vai trò có thể lây lan sang toàn hệ thống."
    )
    document.add_paragraph(
        "Role specialization chỉ có ý nghĩa nếu prompt tạo ra hành vi chức năng khác nhau. Việc đổi nhãn “Critic” thành “Skeptic” nhưng cả hai chỉ viết nhận xét chung không tạo diversity chiến lược. Trong thiết kế này, Critic kiểm tra quan hệ suy diễn, Skeptic chủ động bác bỏ bằng phản ví dụ, Evidence Checker chỉ gắn trạng thái support cho claim, còn Judge so sánh chất lượng chứng cứ. Sự khác nhau được đo bằng nội dung issue/evidence và strategy label, không suy ra từ tên role."
    )

    document.add_heading("1.4. Majority Voting", level=2)
    document.add_paragraph(
        "Majority Voting là quy tắc chọn nhãn có số phiếu lớn nhất. Với N agent hoặc N sample, confidence đồng thuận thường được tính bằng tỷ lệ phiếu của nhãn thắng. Quy tắc này bất biến với thứ tự agent, không cần một model Judge bổ sung và dễ kiểm tra. Trong hệ thống thực tế, đây là ưu điểm lớn: aggregation không phát sinh thêm token, không có position bias ở bước cuối và ít điểm lỗi phần mềm hơn."
    )
    add_equation(document, "ŷMV = arg max(y ∈ C) Σᵢ₌₁ᴺ 𝟙(yᵢ = y)", "(2)")
    document.add_paragraph(
        "Tuy nhiên, tỷ lệ phiếu không phải xác suất đúng. Nếu bốn agent cùng sao chép một assumption sai và một agent thiểu số sử dụng đầy đủ premise, Majority Voting vẫn chọn đáp án sai với consensus 0,8. Vấn đề nghiêm trọng hơn khi lỗi tương quan do cùng model nền. Voting cũng bỏ qua mức confidence riêng, evidence status và chất lượng rationale; nó coi một phiếu vô căn cứ ngang với một phiếu được chứng minh tốt."
    )
    document.add_paragraph(
        "Tie-breaking là một chi tiết kỹ thuật có ảnh hưởng nhưng thường không được báo cáo. Với số agent chẵn, tie có thể được giải bằng thứ tự xuất hiện, confidence hoặc một judge phụ. Mỗi lựa chọn tạo thiên kiến khác nhau. Thiết kế mặc định dùng năm agent để giảm tie; ablation hai agent phải chỉ rõ tie policy. Ngoài Accuracy, cần báo cáo disagreement rate để biết voting thực sự tổng hợp bất đồng hay chỉ xác nhận một đồng thuận đã có từ đầu."
    )

    document.add_heading("1.5. Multi-Agent Debate", level=2)
    document.add_paragraph(
        "Multi-Agent Debate cho phép các agent chia sẻ câu trả lời và phản biện qua nhiều vòng trước khi tổng hợp. Du và cộng sự cho thấy nhiều instance mô hình có thể cải thiện factuality và reasoning thông qua tranh luận [6]. Khan và cộng sự nghiên cứu debate trong bối cảnh tính thuyết phục và tính đúng của câu trả lời, nhấn mạnh quan hệ giữa năng lực debater và khả năng người hoặc mô hình Judge nhận biết sự thật [7]. ChatEval áp dụng multi-agent debate cho đánh giá văn bản, cho thấy debate cũng có thể thay đổi chất lượng của chính bộ đánh giá [8]. Các kết quả này tạo động lực, nhưng không bảo đảm mọi kiến trúc debate đều tốt hơn voting trong cùng ngân sách."
    )
    document.add_paragraph(
        "Một protocol debate tối thiểu gồm proposal, critique, revision và decision. Proposal tạo diversity ban đầu; critique làm lộ điểm yếu; revision biến feedback thành thay đổi có thể đo; decision chọn đáp án cuối. Nếu bỏ revision, protocol chỉ là thảo luận không tác động. Nếu chỉ đo đáp án cuối mà không giữ initial output, không thể tính Correction Rate và Degradation Rate. Nếu Judge nhìn thấy tên agent hoặc thứ tự cố định, kết quả có thể phản ánh identity/position bias thay vì reasoning quality."
    )
    document.add_paragraph(
        "Debate có bốn cơ chế lợi ích tiềm năng. Thứ nhất, error exposure: một agent nhận ra điều mà solver bỏ sót. Thứ hai, hypothesis expansion: Skeptic đưa ra cách hiểu hoặc phản ví dụ mới. Thứ ba, evidence grounding: claim được buộc quay lại context. Thứ tư, selective aggregation: Judge có thể chọn một thiểu số đúng dựa trên chất lượng lý lẽ. Mỗi cơ chế đều có failure mode đối ngẫu: critique sai hướng, diversity nhiễu, evidence label sai và Judge chọn lập luận thuyết phục nhưng không đúng."
    )
    document.add_paragraph(
        "Số vòng debate tạo trade-off rõ ràng. Vòng đầu có thể sửa lỗi hiển nhiên với context tương đối ngắn. Vòng sau đưa thêm cơ hội kiểm tra nhưng transcript dài hơn, chi phí tăng, signal bị pha loãng và conformity pressure mạnh hơn. H4 vì thế không dự đoán tăng tuyến tính; nó dự đoán một điểm bão hòa hoặc suy giảm. Khảo sát r = 0, 1, 2, 3 là đủ để nhận diện hình dạng ban đầu của đường cong mà vẫn giữ phạm vi mini-research."
    )

    document.add_heading("1.6. Độ tin cậy trong NLP Reasoning", level=2)
    document.add_heading("1.6.1. Accuracy và chuyển đổi trạng thái", level=3)
    document.add_paragraph(
        "Accuracy là tỷ lệ dự đoán đúng nhưng không cho biết debate đã tạo thay đổi như thế nào. Hai hệ thống cùng Accuracy có thể khác hoàn toàn: một hệ thống sửa nhiều đáp án sai nhưng đồng thời làm hỏng nhiều đáp án đúng; hệ thống kia hầu như không thay đổi. Correction Rate và Degradation Rate được tính trên hai mẫu số khác nhau, tương ứng tập sai ban đầu và tập đúng ban đầu. Báo cáo đồng thời hai chỉ số ngăn việc chỉ trình bày các successful corrections có lợi."
    )
    add_equation(document, "Correction = |Wrong₀ ∩ CorrectR| / |Wrong₀|", "(3)")
    add_equation(document, "Degradation = |Correct₀ ∩ WrongR| / |Correct₀|", "(4)")

    document.add_heading("1.6.2. Diversity", level=3)
    document.add_paragraph(
        "Answer disagreement chỉ kiểm tra có ít nhất hai nhãn khác nhau trong tập đầu ra. Chỉ số này đơn giản nhưng không phân biệt hai trace chọn cùng đáp án bằng hai chiến lược khác nhau. Semantic diversity dùng biểu diễn vector cho rationale và lấy một trừ trung bình cosine similarity giữa mọi cặp. Sentence-BERT [16] là một lựa chọn embedding cố định; pipeline hiện cũng hỗ trợ TF-IDF như phương án offline tái lập. Khi dùng embedding API, model embedding phải tách khỏi model sinh đáp án và được ghi rõ trong cấu hình."
    )
    add_equation(document, "Dsemantic = 1 − [2 / (N(N−1))] Σᵢ<ⱼ cos(eᵢ, eⱼ)", "(5)")
    document.add_paragraph(
        "Semantic diversity cao không nhất thiết là reasoning diversity có ích. Các trace có thể khác từ vựng nhưng cùng dùng một shortcut; ngược lại, hai lời giải cùng chiến lược có thể diễn đạt khác nhau đáng kể. Vì vậy, một subset cần được gán strategy label như elimination, contradiction, ordering, causal commonsense hoặc lexical association. Mối liên hệ diversity–accuracy nên được phân tích theo bin hoặc regression mô tả, không diễn giải thành quan hệ nhân quả."
    )

    document.add_heading("1.6.3. Calibration", level=3)
    document.add_paragraph(
        "Calibration hỏi liệu nhóm dự đoán confidence xấp xỉ p có đúng khoảng p phần hay không. Guo và cộng sự cho thấy mô hình neural hiện đại có thể bị miscalibration dù Accuracy cao [17]. Brier Score [18] đo sai số bình phương giữa confidence và nhãn đúng–sai; ECE chia confidence thành các bin rồi lấy sai lệch có trọng số giữa confidence trung bình và Accuracy. Trong debate, cần phân biệt self-reported confidence của solver, consensus score của voting và confidence của Judge. Không đại lượng nào tự động là xác suất đã hiệu chuẩn."
    )
    add_equation(document, "Brier = (1/M) Σₘ₌₁ᴹ (pₘ − zₘ)²", "(6)")

    document.add_heading("1.6.4. Chi phí, độ trễ và khả năng tái lập", level=3)
    document.add_paragraph(
        "Một phương pháp inference-time chỉ có ý nghĩa thực tiễn khi hiệu quả được đặt cạnh chi phí. Token usage phải tách input, output và tổng; latency cần mean, median và P95. Đối với protocol tuần tự, latency end-to-end không nhất thiết bằng tổng latency nếu các solver độc lập được chạy song song, vì vậy báo cáo phải ghi rõ concurrency policy. Accuracy trên mỗi 1.000 token và token tăng thêm để đạt thêm một điểm phần trăm Accuracy là hai cách mô tả hiệu suất tài nguyên, nhưng chỉ nên tính khi mẫu số và cấu hình so sánh tương thích."
    )

    document.add_heading("1.7. Các công trình liên quan", level=2)
    add_caption(document, "Bảng 2. So sánh các họ phương pháp inference-time liên quan")
    add_table(
        document,
        ["Phương pháp", "Giao tiếp", "Quyết định", "Ưu điểm", "Rủi ro chính"],
        [
            ("Single Direct", "Không", "Một đầu ra", "Rẻ, nhanh", "Phương sai và lỗi đơn điểm"),
            ("CoT", "Không", "Một đầu ra có rationale", "Lộ bước trung gian", "Hợp lý hóa, verbosity bias"),
            ("Self-Consistency", "Không", "Bỏ phiếu K sample", "Giảm lỗi ngẫu nhiên", "Correlated error; tốn K calls"),
            ("Majority Voting", "Không", "Bỏ phiếu N agent", "Đơn giản; không cần Judge", "Bỏ qua chất lượng evidence"),
            ("Self-Refine/Reflexion", "Tự phản hồi", "Agent tự sửa", "Feedback lặp", "Mù điểm yếu chung của một agent"),
            ("Multi-Agent Debate", "Có", "Vote hoặc Judge", "Expose lỗi; chọn thiểu số đúng", "Conformity; error propagation; Judge bias"),
        ],
        widths=[1550, 1300, 1600, 2300, 2610],
        font_size=8.2,
    )
    document.add_paragraph(
        "Bảng 2 cho thấy khác biệt cốt lõi không nằm ở số lượng đầu ra mà ở việc thông tin có được truyền giữa các tác tử trước quyết định hay không. Self-Consistency và Majority Voting tạo ensemble nhưng giữ độc lập; Self-Refine tạo vòng feedback trong một tác tử; debate tạo feedback chéo. Vì thế, so sánh hợp lý phải giữ ngân sách gần nhau và thay đổi communication protocol hoặc decision protocol một cách có kiểm soát."
    )
    document.add_paragraph(
        "Một nhánh liên quan khác là LLM-as-a-Judge. Judge giúp đánh giá rationale và evidence thay vì chỉ đếm phiếu, nhưng bản thân Judge có thể thiên lệch. Nghiên cứu của Chen và cộng sự phân tích nhiều dạng judgement bias [21]; công trình về position bias cho thấy vị trí câu trả lời có thể ảnh hưởng lựa chọn [22]. Thiết kế blind judge, xáo trộn thứ tự và ẩn danh reasoning ID trong prompt là biện pháp giảm rủi ro, không phải bảo đảm loại bỏ hoàn toàn thiên kiến."
    )

    document.add_heading("1.8. Khoảng trống nghiên cứu", level=2)
    document.add_paragraph(
        "Khoảng trống thứ nhất là fair-compute. Nhiều so sánh trực quan đặt một lần gọi single agent cạnh một protocol gồm nhiều agent, nhiều vòng và một Judge. Kết quả đó hữu ích để biết hiệu năng tối đa của cấu hình, nhưng không trả lời debate có hiệu quả hơn việc dùng cùng số call cho sampling độc lập hay không. Nghiên cứu này tách báo cáo không kiểm soát chi phí khỏi báo cáo khớp số call và khớp token."
    )
    document.add_paragraph(
        "Khoảng trống thứ hai là bản chất của diversity. Tỷ lệ bất đồng đáp án không đủ để chứng minh có đa dạng suy luận; embedding distance cũng có thể bị ảnh hưởng bởi cách diễn đạt. Việc kết hợp answer disagreement, semantic diversity và strategy annotation giúp đưa ra kết luận thận trọng hơn. Nghiên cứu không giả định diversity cao là tốt, mà kiểm tra quan hệ giữa diversity, correction và degradation."
    )
    document.add_paragraph(
        "Khoảng trống thứ ba là tách vai trò của communication và decision. Nếu Debate + Judge thắng Majority Voting, nguyên nhân có thể là transcript giúp sửa reasoning, Judge tốt hơn vote, hoặc cả hai. Việc có cả Debate + Majority và Debate + Judge cho phép đối chiếu hai yếu tố. Evidence-Aware Judge tiếp tục kiểm tra liệu cấu trúc claim–support có tạo lợi ích so với Judge tổng quát hay không."
    )
    document.add_paragraph(
        "Khoảng trống thứ tư là báo cáo failure mode cân bằng. Các case study thường ưu tiên ví dụ debate sửa đúng; harmful revision, minority-correct bị bỏ qua và lỗi Judge ít được lượng hóa. Thiết kế hiện tại khóa taxonomy trước thực nghiệm, lưu initial/final outputs và yêu cầu thống kê cả trường hợp Majority Voting tốt hơn. Đây là nền tảng để trả lời RQ5 mà không chọn lọc hậu nghiệm."
    )


def add_chapter_2(document: Document) -> None:
    heading = document.add_heading("CHƯƠNG 2. PHƯƠNG PHÁP MULTI-AGENT DEBATE ĐỀ XUẤT", level=1)
    heading.paragraph_format.page_break_before = True

    document.add_heading("2.1. Phát biểu bài toán", level=2)
    document.add_paragraph(
        "Xét tập dữ liệu D = {(xₘ, yₘ)} gồm M mẫu. Mỗi x chứa sample_id, dataset, context, question và tập lựa chọn C; y là nhãn đúng đã chuẩn hóa. Một solver S nhận x và trạng thái debate h để sinh đối tượng o = (answer, rationale_summary, evidence, confidence). Một protocol P điều phối N solver cùng tập role chuyên biệt qua R vòng, sau đó một decision rule A trả về ŷ. Mục tiêu thực nghiệm là so sánh các cặp (P, A) theo correctness, reliability và resource cost dưới cấu hình được khóa."
    )
    add_equation(document, "ŷ = A(P(x; N, R, roles, budget))", "(7)")
    document.add_paragraph(
        "Nghiên cứu không tối ưu một hàm utility duy nhất vì trọng số giữa Accuracy, token và latency phụ thuộc ứng dụng. Thay vào đó, các đại lượng được báo cáo riêng và dùng đường cong cost–performance. Khi cần mô tả lựa chọn vận hành, có thể xem utility là Accuracy trừ chi phí có trọng số, nhưng mọi kết luận chính vẫn dựa trên metric nguyên gốc."
    )
    add_equation(document, "U(P) = Accuracy(P) − λt·Tokens(P) − λl·Latency(P)", "(8)")

    document.add_heading("2.2. Nguyên tắc thiết kế", level=2)
    add_bullets(
        document,
        [
            "Independence first: các solver ở vòng 0 không được xem câu trả lời của nhau.",
            "Role separation: mỗi vai trò có tiêu chí kiểm tra và schema riêng, không chỉ khác tên.",
            "Observable rationale only: chỉ dùng rationale_summary do model xuất ra, không suy đoán hidden reasoning.",
            "Evidence before persuasion: claim phải được đối chiếu context trước khi Judge quyết định.",
            "Blind aggregation: ẩn danh và xáo trộn thứ tự candidate trước Judge.",
            "Bounded interaction: số vòng, retry, output tokens và backoff đều hữu hạn.",
            "Traceability: lưu raw response, token, latency, prompt fingerprint, seed và config.",
            "Research honesty: mock/smoke chỉ kiểm tra pipeline; không trộn vào bảng kết quả khoa học.",
        ],
    )

    document.add_heading("2.3. Kiến trúc tổng thể", level=2)
    document.add_paragraph(
        "Kiến trúc gồm tám khối theo Hình 1. Câu hỏi được chuẩn hóa thành schema thống nhất. N solver sinh proposal độc lập. Disagreement Detector xác định mức bất đồng ban đầu và quyết định có cần mở debate theo policy hay không. Ba role Critic, Skeptic và Evidence Checker đọc trạng thái được phép, tạo feedback có cấu trúc. Solver nhận feedback để revision. Chu trình lặp R vòng, sau đó Majority Aggregator hoặc Judge tạo quyết định cuối. Tất cả intermediate output được ghi JSONL cùng usage metadata để phục vụ phân tích lại."
    )
    add_architecture_figure(document)
    document.add_paragraph(
        "Việc tách rõ aggregation khỏi debate là chủ ý phương pháp. Cùng một transcript sau revision có thể đi qua Majority Vote hoặc Judge; chênh lệch giữa hai kết quả phản ánh decision protocol. Ngược lại, so sánh independent Majority Voting với Debate + Majority giữ cơ chế quyết định gần như cố định, từ đó cô lập phần đóng góp của communication. Evidence-Aware Judge là biến thể mạnh hơn, nhưng phải được báo cáo như một cấu hình riêng vì nó dùng thêm cấu trúc evidence."
    )

    document.add_heading("2.4. Biểu diễn đầu vào và chuẩn hóa đáp án", level=2)
    document.add_paragraph(
        "Canonical Sample gồm sample_id duy nhất, dataset, split, context, question, choices và answer. Nhãn đáp án được chuẩn hóa về ký tự in hoa đầu tiên thuộc tập lựa chọn. Những mẫu thiếu đáp án, không ánh xạ được lựa chọn hoặc có cấu trúc không hợp lệ bị loại trước khi khóa subset. Danh sách ID được lưu độc lập để mọi seed và mọi phương pháp đánh giá cùng một tập mẫu."
    )
    document.add_paragraph(
        "Prompt user ghép metadata tối thiểu với nội dung câu hỏi và trạng thái debate. Metadata như sample_id, dataset và round hỗ trợ truy vết nhưng không được chứa đáp án chuẩn. Khi truyền transcript, hệ thống ưu tiên JSON đã cấu trúc thay vì chuỗi tự do. Context được giới hạn theo ngân sách; nếu vượt ngưỡng, policy cần cắt ở ranh giới message hoặc tóm tắt có lưu provenance, không cắt giữa một đối tượng JSON."
    )

    document.add_heading("2.5. Thiết kế các vai trò", level=2)
    document.add_heading("2.5.1. Solver", level=3)
    document.add_paragraph(
        "Solver phân tích câu hỏi, chọn một đáp án, tạo rationale_summary ngắn, liệt kê evidence và tự báo confidence. Ở vòng 0, mỗi solver chạy độc lập với cùng task prompt nhưng khác sampling path/agent index. Ở vòng r > 0, solver nhận previous_answer và critiques, sau đó phải quyết định sửa hoặc bảo lưu. Việc bảo lưu được cho phép để tránh prompt mặc nhiên gợi rằng critique chắc chắn đúng."
    )
    add_bullets(
        document,
        [
            "Không thêm dữ kiện ngoài context nếu bài toán yêu cầu reasoning nội tại.",
            "Mỗi rationale item trình bày một quan hệ suy luận kiểm tra được.",
            "Mỗi evidence item nêu claim, source và status.",
            "Answer phải thuộc nhãn lựa chọn hợp lệ; confidence bị chặn trong [0, 1].",
            "Revision phải phản hồi critique theo nội dung, không đổi đáp án chỉ để đạt đồng thuận.",
        ],
    )

    document.add_heading("2.5.2. Critic", level=3)
    document.add_paragraph(
        "Critic kiểm tra từng khẳng định trong rationale, tập trung vào invalid implication, contradiction, missing premise và assumption không được hỗ trợ. Mỗi issue có type, target_step, description và severity. Critic không được dùng nhận xét chung như “hãy suy nghĩ kỹ hơn,” cũng không được mặc định solver sai. Nếu không phát hiện lỗi đáng kể, output issues có thể rỗng; hành vi này quan trọng để giảm áp lực sửa không cần thiết."
    )

    document.add_heading("2.5.3. Skeptic", level=3)
    document.add_paragraph(
        "Skeptic chủ động thử bác bỏ kết luận hiện tại. Vai trò này tìm phản ví dụ, cách hiểu thay thế, trường hợp biên hoặc phương án đối thủ có vẻ hợp lý. Khác Critic, Skeptic không chỉ audit từng bước đã có mà mở rộng không gian giả thuyết. Prompt yêu cầu tránh lặp nguyên văn Critic; trong phân tích, mức trùng lặp issue giữa hai role có thể được dùng để kiểm tra liệu specialization có tạo hành vi khác biệt thật hay không."
    )

    document.add_heading("2.5.4. Evidence Checker", level=3)
    document.add_paragraph(
        "Evidence Checker đối chiếu claim với context và điều kiện logic. Mỗi claim nhận một trong bốn nhãn: SUPPORTED nếu được đề bài hoặc phép suy diễn hợp lệ hỗ trợ; UNSUPPORTED nếu không có căn cứ đủ; CONTRADICTED nếu xung đột với premise; UNCERTAIN nếu dữ liệu không cho phép kết luận. Vai trò này không thay thế solver và không nên tự thêm kiến thức mới; nhiệm vụ chính là tạo lớp grounding có cấu trúc cho revision và Judge."
    )

    document.add_heading("2.5.5. Revision", level=3)
    document.add_paragraph(
        "Revision được thực hiện bởi solver trên cùng schema SolverOutput, nhờ đó đáp án vòng 0 và vòng sau có thể so sánh trực tiếp. Mỗi reasoning_id mã hóa solver index và round. Hệ thống giữ cả initial_outputs và revised outputs; không ghi đè bản trước. Thiết kế này cho phép tính correction/degradation, truy tác nhân phát hiện lỗi và nhận diện conformity khi solver đổi theo số đông dù evidence không cải thiện."
    )

    document.add_heading("2.5.6. Judge", level=3)
    document.add_paragraph(
        "Judge so sánh candidate dựa trên reasoning quality, evidence support và consistency với câu hỏi. Trước khi gọi Judge, candidate được xáo trộn bằng seed cố định và thay định danh có ý nghĩa bằng reasoning_id ẩn danh. Judge không nhận gold answer, không được ưu tiên câu trả lời dài hơn và không chỉ đếm phiếu. Output gồm final_answer, selected_reasoning_id, decision_reason và confidence."
    )
    document.add_paragraph(
        "Evidence-Aware Judge nhận thêm evidence status và ưu tiên candidate có claim được hỗ trợ, nhưng vẫn phải tự kiểm tra tính liên quan. Majority Judge không phải cấu hình hợp lệ trong nghiên cứu vì sẽ trộn hai protocol; Majority Voting được triển khai bằng code xác định, còn Judge là một model call riêng có usage và latency được ghi nhận."
    )

    document.add_heading("2.6. Disagreement Detector", level=2)
    document.add_paragraph(
        "Disagreement Detector tạo các đặc trưng trước debate: số nhãn khác nhau, tỷ lệ phiếu lớn nhất, entropy của phân phối phiếu, độ phân tán confidence và semantic diversity giữa rationale. Điều kiện bất đồng tối thiểu là tồn tại ít nhất hai đáp án khác nhau. Với p(c) là tỷ lệ agent chọn c, normalized entropy có thể dùng để so sánh giữa bài có số lựa chọn khác nhau."
    )
    add_equation(document, "Hnorm = −Σc∈C p(c) log p(c) / log |C|", "(9)")
    document.add_paragraph(
        "Trong main experiment, debate rounds được chạy theo cấu hình cố định để so sánh trực tiếp. Adaptive debate là phân tích mở rộng: chỉ kích hoạt vòng tiếp theo khi disagreement cao, evidence còn UNSUPPORTED/CONTRADICTED hoặc confidence chưa ổn định. Adaptive policy có thể giảm chi phí nhưng tạo selection effect, nên không được trộn với kết quả fixed-round nếu không báo cáo riêng."
    )

    document.add_heading("2.7. Debate protocol và điều kiện dừng", level=2)
    document.add_paragraph(
        "Vòng 0 tạo N SolverOutput độc lập. Ở mỗi vòng r từ 1 đến R, các role chuyên biệt đọc trạng thái ở vòng trước và tạo feedback. Sau đó từng solver nhận cùng tập feedback nhưng giữ previous_answer riêng để revision. State của vòng mới gồm revised outputs và critiques vừa tạo. Hệ thống không cho Judge can thiệp giữa vòng nhằm tránh feedback bị định hướng bởi sở thích của Judge."
    )
    document.add_paragraph(
        "Fixed-round policy dừng sau R ∈ {0,1,2,3}. Một optional early-stop policy có thể dừng nếu: tất cả solver cùng đáp án trong hai trạng thái liên tiếp; không có issue severity HIGH; không còn claim CONTRADICTED; và mức thay đổi confidence dưới ngưỡng ε. Các tiêu chí phải được cấu hình, không hard-code, và trạng thái dừng phải được log. Để ablation số vòng không bị nhiễu, kết quả chính ưu tiên fixed rounds."
    )
    document.add_heading("2.7.1. Thuật toán tổng quát", level=3)
    algorithm = [
        "Algorithm 1: Role-Specialized Multi-Agent Debate",
        "Input: sample x, number of solvers N, rounds R, roles Q, decision rule A, seed s",
        "1. O0 ← [Solver(x, round=0, agent_index=i)] for i = 1..N",
        "2. state ← {initial: O0}; save(O0)",
        "3. for r = 1..R:",
        "4.     F ← []",
        "5.     for role q in {Critic, Skeptic, EvidenceChecker} minus removed_roles:",
        "6.         F.append(q(x, round=r, state))",
        "7.     Or ← [Solver.revise(x, previous=Or−1[i], feedback=F)] for i = 1..N",
        "8.     state ← {previous: Or, critiques: F}; save(F, Or)",
        "9. candidates ← shuffle_and_anonymize(OR, seed=s)",
        "10. ŷ ← MajorityVote(candidates) if A=majority else Judge(x, candidates, A)",
        "11. save(ŷ, token_usage, latency, configuration_hash)",
        "Output: final prediction ŷ and complete auditable trace",
    ]
    for line in algorithm:
        p = document.add_paragraph(style="Code Block")
        p.add_run(line)

    document.add_heading("2.8. Giao thức quyết định", level=2)
    document.add_heading("2.8.1. Debate + Majority Voting", level=3)
    document.add_paragraph(
        "Sau vòng cuối, hệ thống chỉ đếm answer của N solver. Confidence được ghi như consensus score = số phiếu nhãn thắng/N và phải mang đúng nhãn loại confidence trong dữ liệu xử lý. Cấu hình này kiểm tra tác động của giao tiếp khi decision rule vẫn đơn giản. Nếu số phiếu hòa, policy cần dùng tie-break xác định trước; lựa chọn mặc định ưu tiên số agent lẻ."
    )

    document.add_heading("2.8.2. Debate + Blind Judge", level=3)
    document.add_paragraph(
        "Judge nhận candidate đã xáo trộn và ẩn danh. Seed kiểm soát hoán vị, vì position bias có thể tồn tại ngay cả khi prompt yêu cầu công bằng [22]. Để kiểm tra độ nhạy, có thể lặp Judge với nhiều permutation trên subset nhỏ, nhưng các lượt này phải tính vào compute. Judge output được validate như mọi agent khác và một lần retry sửa định dạng không được che giấu trong log."
    )

    document.add_heading("2.8.3. Debate + Evidence-Aware Judge", level=3)
    document.add_paragraph(
        "Evidence-Aware Judge dùng cùng candidate nhưng có rubric tường minh: coverage của premise, số claim SUPPORTED, sự hiện diện của CONTRADICTED/UNSUPPORTED, tính nhất quán giữa answer và rationale, và mức độ xử lý critique. Rubric không biến evidence count thành luật cứng; một claim trọng yếu có thể quan trọng hơn nhiều claim phụ. Cấu hình này nhằm kiểm tra liệu cấu trúc evidence mang lại giá trị ngoài khả năng đánh giá tự do của Judge."
    )

    document.add_heading("2.9. Hợp đồng JSON và xác thực", level=2)
    document.add_paragraph(
        "Mọi output dùng JSON hợp lệ để loại bỏ phụ thuộc vào parser chuỗi tự do. Pydantic xác thực kiểu trường, enum và khoảng confidence. Field name giữ tiếng Anh để ổn định code; nội dung rationale, description, decision_reason và recommended_revision dùng tiếng Việt có dấu. Nếu response chứa wrapper Markdown, parser chỉ được phép tách một đối tượng JSON rõ ràng; nếu schema sai, hệ thống retry hữu hạn với thông báo sửa định dạng và ghi lại lần thất bại."
    )
    add_caption(document, "Bảng 3. Hợp đồng dữ liệu của các agent")
    add_table(
        document,
        ["Role", "Trường bắt buộc chính", "Ràng buộc quan trọng"],
        [
            ("Solver/Revision", "sample_id, agent_role, round, answer, rationale_summary, evidence, confidence", "answer hợp lệ; confidence ∈ [0,1]; evidence có status"),
            ("Critic/Skeptic", "sample_id, agent_role, round, issues, recommended_revision", "issue type thuộc taxonomy; severity LOW/MEDIUM/HIGH"),
            ("Evidence Checker", "sample_id, agent_role, round, evidence, recommended_revision", "status thuộc bốn nhãn support"),
            ("Judge", "sample_id, final_answer, selected_reasoning_id, decision_reason, confidence", "không có gold; answer hợp lệ; confidence ∈ [0,1]"),
        ],
        widths=[1500, 4700, 3160],
        font_size=8.5,
    )
    solver_json = [
        '{',
        '  "sample_id": "logiqa_001",',
        '  "agent_role": "solver",',
        '  "round": 0,',
        '  "answer": "B",',
        '  "rationale_summary": ["Tiền đề 1 loại trừ A.", "Điều kiện 3 suy ra B."],',
        '  "evidence": [{"claim": "A bị loại", "source": "premise_1", "status": "SUPPORTED"}],',
        '  "confidence": 0.78',
        '}',
    ]
    for line in solver_json:
        document.add_paragraph(line, style="Code Block")

    document.add_heading("2.10. Retry, rate limit, caching và logging", level=2)
    document.add_paragraph(
        "LLM client đọc API key từ biến môi trường và model name từ YAML. Mỗi request được băm từ model, system prompt, user prompt, role, temperature và max output tokens. Nếu cache có khóa tương ứng, response được tái sử dụng và gắn cached=true. Nếu gọi API thật, client áp dụng min delay, retry tối đa và backoff tăng theo lần thử. Mọi giới hạn đều nằm trong cấu hình để có thể điều chỉnh theo provider."
    )
    document.add_paragraph(
        "Cache dạng JSONL hỗ trợ append-only và lưu content, usage, latency, model, role. Structured log bổ sung timestamp UTC, level, module, sample_id, method, seed và trạng thái checkpoint. Raw prediction record tách answer/gold/correct khỏi raw agent outputs, giúp evaluator hoạt động mà không cần gọi model. Resume bỏ qua các khóa sample–method–seed đã hoàn thành; overwrite mặc định false để tránh mất dữ liệu."
    )
    document.add_paragraph(
        "Một giới hạn cần ghi rõ là retry do network và retry do schema invalid không tương đương. Network retry thường không tạo completion tính phí nếu request thất bại trước inference, trong khi schema retry có thể phát sinh token mới. Báo cáo chi phí phải lấy usage thực tế của mọi completion đã nhận, kể cả completion cuối cùng không dùng nếu nó đã được provider tính phí và được lưu trong log."
    )

    document.add_heading("2.11. Mô hình chi phí và fair-compute", level=2)
    document.add_paragraph(
        "Với N solver, R vòng, Q role phản biện hoạt động ở mỗi vòng và một Judge tùy chọn, số model call tuần tự theo đặc tả là N + R(N + Q) + J, trong đó J ∈ {0,1}. Với cấu hình mặc định N=5, Q=3, một vòng và Judge, một mẫu cần 14 calls; ba vòng cần 30 calls. Đây là lý do không thể so sánh trực tiếp với Single Direct một call rồi quy toàn bộ chênh lệch Accuracy cho debate."
    )
    add_equation(document, "CallsMAD = N + R(N + Q) + J", "(10)")
    document.add_paragraph(
        "Fair-compute theo call chọn K của Self-Consistency hoặc số candidate của Majority Voting sao cho tổng call gần CallsMAD. Fair-compute theo token khó hơn vì prompt critique và Judge dài hơn proposal. Một scheduler token-aware cần theo dõi cumulative input/output tokens và dừng khi chạm budget đã xác định trước. Nếu hai phương pháp vượt budget khác nhau, báo cáo cả intended budget, realized tokens và độ lệch."
    )
    document.add_paragraph(
        "Hai chỉ số hiệu suất phụ được định nghĩa là Accuracy trên mỗi 1.000 token và token tăng thêm để đạt thêm một điểm phần trăm Accuracy. Chỉ số thứ hai không xác định khi Accuracy không tăng hoặc giảm; trong trường hợp đó báo cáo “không có lợi ích Accuracy” thay vì số âm gây hiểu nhầm. Không dùng giá tiền API làm đơn vị duy nhất vì pricing có thể thay đổi theo thời gian và provider."
    )

    document.add_heading("2.12. Cơ chế hạn chế failure mode", level=2)
    add_caption(document, "Bảng 4. Failure mode và biện pháp kiểm soát trong thiết kế")
    add_table(
        document,
        ["Failure mode", "Biểu hiện", "Biện pháp thiết kế", "Tín hiệu đánh giá"],
        [
            ("Majority pressure", "Solver đổi theo số đông dù evidence yếu", "Giữ previous answer; cho phép bảo lưu; blind candidate", "Correct→Wrong; change không tăng support"),
            ("Sycophancy", "Phản hồi đồng ý critique thiếu căn cứ", "Critique không được mặc định đúng; yêu cầu lý do revision", "Tỷ lệ đổi theo critique sai"),
            ("Error propagation", "Một claim sai lan qua nhiều vòng", "Evidence status; lưu provenance; giới hạn rounds", "UNSUPPORTED/CONTRADICTED tái xuất hiện"),
            ("Judge bias", "Ưu tiên vị trí, độ dài hoặc phong cách", "Shuffle seed; ẩn danh; rubric; permutation audit", "Flip theo hoán vị; Judge error"),
            ("Correlated error", "Nhiều agent cùng shortcut sai", "Role specialization; diversity measurement", "Consensus cao nhưng sai"),
            ("Context overload", "Transcript dài làm bỏ sót premise", "Bounded rounds/tokens; state có cấu trúc", "Accuracy/latency giảm ở vòng sâu"),
            ("Format failure", "JSON sai hoặc answer không hợp lệ", "Pydantic; retry hữu hạn; raw log", "Validation/retry rate"),
        ],
        widths=[1350, 2350, 3000, 2660],
        font_size=8.0,
    )
    document.add_paragraph(
        "Các biện pháp trên giảm rủi ro nhưng không loại bỏ chúng. Ví dụ, xáo trộn một lần không chứng minh Judge không có position bias; Evidence Checker cũng có thể gắn nhãn sai; role specialization có thể chỉ thay đổi phong cách. Vì vậy, mỗi biện pháp luôn đi cùng tín hiệu đánh giá có thể bác bỏ kỳ vọng. Đây là nguyên tắc falsifiability ở cấp thiết kế hệ thống."
    )

    document.add_heading("2.13. Liên hệ giữa đặc tả và repository", level=2)
    add_table(
        document,
        ["Thành phần phương pháp", "Vị trí triển khai", "Artifact đầu ra"],
        [
            ("Agent wrapper và validation", "src/agents/base.py; src/schemas/agent_outputs.py", "SolverOutput, CritiqueOutput, EvidenceCheckerOutput, JudgeOutput"),
            ("OpenAI client, retry và usage", "src/agents/llm_client.py", "content, input/output tokens, latency, cached"),
            ("Baseline", "src/protocols/baselines.py", "PredictionRecord cho single, SC và majority"),
            ("Debate orchestration", "src/protocols/debate.py", "initial outputs, critiques, revisions, final output"),
            ("Decision protocol", "src/protocols/aggregation.py", "majority result hoặc judge-converted result"),
            ("Metrics", "src/evaluation/metrics.py", "accuracy, diversity, calibration, CI, McNemar"),
            ("Cache/log", "src/utils/cache.py; src/utils/logging.py", "response_cache.jsonl; structured logs"),
        ],
        widths=[2450, 3250, 3660],
        font_size=8.2,
    )
    document.add_paragraph(
        "Bảng ánh xạ này phản ánh prototype hiện tại và là căn cứ kiểm tra tính nhất quán giữa báo cáo với code. Một số khả năng như adaptive early stopping, permutation audit của Judge và embedding API được mô tả như phần mở rộng có cấu hình; chúng không được tuyên bố là đã có kết quả. Chương 3 sau này phải ghi chính xác phiên bản code/config đã dùng cho mỗi run."
    )


def add_pending_sections(document: Document) -> None:
    heading = document.add_heading("CHƯƠNG 3. THIẾT KẾ THỰC NGHIỆM VÀ ĐÁNH GIÁ", level=1)
    heading.paragraph_format.page_break_before = True
    add_note(
        document,
        "CHƯA HOÀN THIỆN. Chương này sẽ được viết sau khi pilot khóa prompt, parser, subset, model configuration và ngân sách. Nội dung tương lai gồm dataset, split, model, baseline, main experiment, fair-compute, ablation, metrics, statistical tests và reproducibility protocol.",
    )

    document.add_heading("CHƯƠNG 4. KẾT QUẢ VÀ THẢO LUẬN", level=1)
    add_note(
        document,
        "CHƯA HOÀN THIỆN. Chương này chỉ được viết từ raw output API thật và bảng sinh tự động sau main experiment/ablation. Không sử dụng số liệu mock, dry-run hoặc smoke test để kiểm định RQ và hypotheses.",
    )

    document.add_heading("KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", level=1)
    document.add_heading("1. Kết luận", level=2)
    document.add_paragraph("Sẽ hoàn thiện sau khi có kết quả thực nghiệm thật; bản thảo hiện tại không kết luận MAD tốt hơn hoặc kém hơn Majority Voting.")
    document.add_heading("2. Hạn chế", level=2)
    document.add_paragraph("Sẽ tổng hợp từ giới hạn dữ liệu, model, prompt, compute, annotation và kiểm định sau khi hoàn tất pipeline thực nghiệm.")
    document.add_heading("3. Hướng phát triển", level=2)
    document.add_paragraph("Sẽ ưu tiên các hướng được gợi ý trực tiếp bởi error analysis và cost–performance analysis, thay vì liệt kê mở rộng không có căn cứ thực nghiệm.")


REFERENCES = [
    ("[1]", "A. Vaswani et al., “Attention Is All You Need,” in Advances in Neural Information Processing Systems, vol. 30, 2017.", "https://proceedings.neurips.cc/paper/2017/hash/3f5ee243547dee91fbd053c1c4a845aa-Abstract.html"),
    ("[2]", "T. B. Brown et al., “Language Models are Few-Shot Learners,” in Advances in Neural Information Processing Systems, vol. 33, pp. 1877–1901, 2020.", "https://proceedings.neurips.cc/paper/2020/hash/1457c0d6bfcb4967418bfb8ac142f64a-Abstract.html"),
    ("[3]", "J. Wei et al., “Chain-of-Thought Prompting Elicits Reasoning in Large Language Models,” in Advances in Neural Information Processing Systems, vol. 35, pp. 24824–24837, 2022.", "https://proceedings.neurips.cc/paper_files/paper/2022/hash/9d5609613524ecf4f15af0f7b31abca4-Abstract-Conference.html"),
    ("[4]", "T. Kojima, S. S. Gu, M. Reid, Y. Matsuo, and Y. Iwasawa, “Large Language Models are Zero-Shot Reasoners,” in Advances in Neural Information Processing Systems, vol. 35, 2022.", "https://proceedings.neurips.cc/paper_files/paper/2022/hash/8bb0d291acd4acf06ef112099c16f326-Abstract-Conference.html"),
    ("[5]", "X. Wang et al., “Self-Consistency Improves Chain of Thought Reasoning in Language Models,” in Proc. International Conference on Learning Representations, 2023.", "https://openreview.net/forum?id=1PL1NIMMrw"),
    ("[6]", "Y. Du, S. Li, A. Torralba, J. B. Tenenbaum, and I. Mordatch, “Improving Factuality and Reasoning in Language Models through Multiagent Debate,” in Proc. 41st International Conference on Machine Learning, PMLR 235, pp. 11733–11763, 2024.", "https://proceedings.mlr.press/v235/du24e.html"),
    ("[7]", "A. Khan et al., “Debating with More Persuasive LLMs Leads to More Truthful Answers,” in Proc. 41st International Conference on Machine Learning, PMLR 235, pp. 23662–23733, 2024.", "https://proceedings.mlr.press/v235/khan24a.html"),
    ("[8]", "C.-M. Chan et al., “ChatEval: Towards Better LLM-based Evaluators through Multi-Agent Debate,” in Proc. International Conference on Learning Representations, 2024.", "https://openreview.net/forum?id=FQepisCUWu"),
    ("[9]", "A. Madaan et al., “Self-Refine: Iterative Refinement with Self-Feedback,” in Advances in Neural Information Processing Systems, vol. 36, 2023.", "https://proceedings.neurips.cc/paper_files/paper/2023/hash/91edff07232fb1b55a505a9e9f6c0ff3-Abstract-Conference.html"),
    ("[10]", "N. Shinn, F. Cassano, A. Gopinath, K. Narasimhan, and S. Yao, “Reflexion: Language Agents with Verbal Reinforcement Learning,” in Advances in Neural Information Processing Systems, vol. 36, 2023.", "https://proceedings.neurips.cc/paper_files/paper/2023/hash/1b44b878bb782e6954cd888628510e90-Abstract-Conference.html"),
    ("[11]", "S. Yao et al., “ReAct: Synergizing Reasoning and Acting in Language Models,” in Proc. International Conference on Learning Representations, 2023.", "https://openreview.net/forum?id=WE_vluYUL-X"),
    ("[12]", "J. Liu, L. Cui, H. Liu, D. Huang, Y. Wang, and Y. Zhang, “LogiQA: A Challenge Dataset for Machine Reading Comprehension with Logical Reasoning,” in Proc. 29th International Joint Conference on Artificial Intelligence, pp. 3622–3628, 2020, doi: 10.24963/ijcai.2020/501.", "https://www.ijcai.org/proceedings/2020/501"),
    ("[13]", "A. Talmor, J. Herzig, N. Lourie, and J. Berant, “CommonsenseQA: A Question Answering Challenge Targeting Commonsense Knowledge,” in Proc. NAACL-HLT, pp. 4149–4158, 2019, doi: 10.18653/v1/N19-1421.", "https://aclanthology.org/N19-1421/"),
    ("[14]", "M. Geva, D. Khashabi, E. Segal, T. Khot, D. Roth, and J. Berant, “Did Aristotle Use a Laptop? A Question Answering Benchmark with Implicit Reasoning Strategies,” Transactions of the Association for Computational Linguistics, vol. 9, pp. 346–361, 2021, doi: 10.1162/tacl_a_00370.", "https://aclanthology.org/2021.tacl-1.21/"),
    ("[15]", "K. Cobbe et al., “Training Verifiers to Solve Math Word Problems,” arXiv:2110.14168, 2021.", "https://arxiv.org/abs/2110.14168"),
    ("[16]", "N. Reimers and I. Gurevych, “Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks,” in Proc. EMNLP-IJCNLP, pp. 3982–3992, 2019, doi: 10.18653/v1/D19-1410.", "https://aclanthology.org/D19-1410/"),
    ("[17]", "C. Guo, G. Pleiss, Y. Sun, and K. Q. Weinberger, “On Calibration of Modern Neural Networks,” in Proc. 34th International Conference on Machine Learning, PMLR 70, pp. 1321–1330, 2017.", "https://proceedings.mlr.press/v70/guo17a.html"),
    ("[18]", "G. W. Brier, “Verification of Forecasts Expressed in Terms of Probability,” Monthly Weather Review, vol. 78, no. 1, pp. 1–3, 1950, doi: 10.1175/1520-0493(1950)078<0001:VOFEIT>2.0.CO;2.", "https://journals.ametsoc.org/view/journals/mwre/78/1/1520-0493_1950_078_0001_vofeit_2_0_co_2.xml"),
    ("[19]", "Q. McNemar, “Note on the Sampling Error of the Difference Between Correlated Proportions or Percentages,” Psychometrika, vol. 12, pp. 153–157, 1947, doi: 10.1007/BF02295996.", "https://doi.org/10.1007/BF02295996"),
    ("[20]", "B. Efron, “Bootstrap Methods: Another Look at the Jackknife,” The Annals of Statistics, vol. 7, no. 1, pp. 1–26, 1979, doi: 10.1214/aos/1176344552.", "https://doi.org/10.1214/aos/1176344552"),
    ("[21]", "G. Chen et al., “Humans or LLMs as the Judge? A Study on Judgement Bias,” in Proc. 2024 Conference on Empirical Methods in Natural Language Processing, pp. 8301–8327, 2024, doi: 10.18653/v1/2024.emnlp-main.474.", "https://aclanthology.org/2024.emnlp-main.474/"),
    ("[22]", "L. Shi, J. Li, S. Zhang, Z. Lai, and H. Shen, “Judging the Judges: A Systematic Study of Position Bias in LLM-as-a-Judge,” in Proc. IJCNLP-AACL, pp. 292–314, 2025, doi: 10.18653/v1/2025.ijcnlp-long.18.", "https://aclanthology.org/2025.ijcnlp-long.18/"),
    ("[23]", "S. Lin, J. Hilton, and O. Evans, “TruthfulQA: Measuring How Models Mimic Human Falsehoods,” in Proc. 60th Annual Meeting of the Association for Computational Linguistics, pp. 3214–3252, 2022, doi: 10.18653/v1/2022.acl-long.229.", "https://aclanthology.org/2022.acl-long.229/"),
    ("[24]", "J. Cohen, “A Coefficient of Agreement for Nominal Scales,” Educational and Psychological Measurement, vol. 20, no. 1, pp. 37–46, 1960, doi: 10.1177/001316446002000104.", "https://doi.org/10.1177/001316446002000104"),
]


def add_references_and_appendix(document: Document) -> None:
    heading = document.add_heading("TÀI LIỆU THAM KHẢO", level=1)
    heading.paragraph_format.page_break_before = True
    for label, citation, url in REFERENCES:
        p = document.add_paragraph(style="Reference")
        p.add_run(f"{label} {citation} ")
        add_hyperlink(p, "[Trực tuyến]", url)

    heading = document.add_heading("PHỤ LỤC", level=1)
    heading.paragraph_format.page_break_before = True
    document.add_heading("Phụ lục A. Checklist trung thực nghiên cứu", level=2)
    add_table(
        document,
        ["Nguyên tắc", "Cách thực thi trong dự án"],
        [
            ("Không bịa số liệu", "Chương 4 để trống cho đến khi có raw JSONL từ API thật."),
            ("Không bịa citation", "Tài liệu tham khảo dùng paper/proceedings chính thức và DOI/URL kiểm chứng."),
            ("Không prompt-tune trên test", "Prompt phát triển trên dev; sample ID của test subset khóa trước main run."),
            ("Không đồng nhất consensus với confidence", "Lưu consensus_score, solver_confidence và judge_confidence như các trường khác nhau."),
            ("Không chọn case có lợi", "Behavioral taxonomy bao gồm correction, resistant error, harmful revision và Majority-wins."),
            ("Không bỏ qua compute", "Báo cáo riêng raw Accuracy, call-matched và token-matched comparisons."),
        ],
        widths=[2500, 6860],
    )

    document.add_heading("Phụ lục B. Taxonomy lỗi khóa trước thực nghiệm", level=2)
    add_table(
        document,
        ["Mã lỗi", "Định nghĩa thao tác"],
        [
            ("LOGICAL_ERROR", "Kết luận không suy ra từ tiền đề hoặc vi phạm ràng buộc logic."),
            ("MISINTERPRETATION", "Hiểu sai câu hỏi, phạm vi lượng từ, phủ định hoặc ý nghĩa lựa chọn."),
            ("MISSING_EVIDENCE", "Bỏ sót dữ kiện cần thiết trong context."),
            ("UNSUPPORTED_ASSUMPTION", "Thêm giả định không có căn cứ đủ."),
            ("ARITHMETIC_ERROR", "Sai phép tính hoặc thao tác định lượng."),
            ("HALLUCINATION", "Tạo dữ kiện, nguồn hoặc quan hệ không tồn tại."),
            ("CONFORMITY_ERROR", "Đổi từ đúng sang sai do áp lực đồng thuận/feedback."),
            ("JUDGE_ERROR", "Judge chọn candidate kém hơn dù candidate đúng có bằng chứng đủ."),
            ("ANSWER_EXTRACTION_ERROR", "Reasoning hướng đúng nhưng nhãn output không khớp hoặc parser sai."),
            ("CONTEXT_OVERLOAD", "Bỏ sót/nhầm thông tin do transcript hoặc context quá dài."),
        ],
        widths=[2750, 6610],
    )

    document.add_heading("Phụ lục C. Trạng thái hoàn thiện", level=2)
    add_table(
        document,
        ["Phần", "Trạng thái", "Điều kiện hoàn thiện tiếp"],
        [
            ("Mở đầu", "Đã viết", "Cập nhật phạm vi nếu cấu hình main thay đổi."),
            ("Chương 1", "Đã viết", "Chỉ bổ sung tài liệu mới nếu trực tiếp liên quan kết quả."),
            ("Chương 2", "Đã viết", "Đồng bộ nếu code/protocol thay đổi trước khóa phiên bản."),
            ("Chương 3", "Chờ thực nghiệm", "Hoàn tất pilot, khóa prompt, subset, model và budget."),
            ("Chương 4", "Chờ thực nghiệm", "Có raw output thật, evaluation tables, figures và statistical tests."),
            ("Kết luận", "Chờ thực nghiệm", "Trả lời RQ1–RQ5 và phân loại H1–H5 từ bằng chứng."),
        ],
        widths=[1900, 1900, 5560],
    )


def set_document_properties(document: Document) -> None:
    props = document.core_properties
    props.title = "Beyond Majority Voting: Multi-Agent Debate for Reliable NLP Reasoning"
    props.subject = "Báo cáo tiểu luận khoa học – Mở đầu, Chương 1 và Chương 2"
    props.author = "Nhóm nghiên cứu NLP"
    props.keywords = "Multi-Agent Debate, Majority Voting, NLP Reasoning, LLM, Reliability"
    props.comments = "Không chứa kết quả thực nghiệm giả; Chương 3 và Chương 4 chờ main experiment."

    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def build_report() -> Path:
    document = Document()
    configure_styles(document)
    set_document_properties(document)
    add_cover(document)
    add_front_matter(document)
    add_introduction(document)
    add_chapter_1(document)
    add_chapter_2(document)
    add_pending_sections(document)
    add_references_and_appendix(document)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    path = build_report()
    print(path)
